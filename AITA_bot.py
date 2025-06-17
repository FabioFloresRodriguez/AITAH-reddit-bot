import praw
import re
import docx
from pathlib import Path
from datetime import date
from original_title import Original_Title

reddit = praw.Reddit("bot1")

docx_path = Path(f"Weekly AITAH stories/{str(date.today())} AITAH stories.docx")

doc = docx.Document()
doc.save(docx_path)

subreddit = reddit.subreddit("AITAH")
for submission in subreddit.hot(limit=11):
    if re.search("AITAH for banning users with scam links and other domains mostly bots use?",submission.title,re.IGNORECASE):
        continue

    title = submission.title
    author = reddit.redditor(f"{submission.author}")
    full_story = {}
    if re.search("update",submission.title,re.IGNORECASE):
        title = Original_Title(submission.title)

    for author_submission in author.submissions.new(limit=None):
        if re.search(title,author_submission.title,re.IGNORECASE) and not re.search("update",author_submission.title,re.IGNORECASE):
            full_story[author_submission.title] = author_submission.selftext 
        elif re.search(title,author_submission.title,re.IGNORECASE) and re.search("update",author_submission.title,re.IGNORECASE) and not re.search("final",author_submission.title,re.IGNORECASE):
            full_story[author_submission.title] = author_submission.selftext
        elif re.search(title,author_submission.title,re.IGNORECASE) and re.search("final",author_submission.title,re.IGNORECASE):
            full_story[author_submission.title] = author_submission.selftext

    full_story = dict(reversed(list(full_story.items())))

    for key,val in full_story.items():
        doc.add_heading(key,1)
        doc.add_paragraph(val)

doc.save(docx_path)